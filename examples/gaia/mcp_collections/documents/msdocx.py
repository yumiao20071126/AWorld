import json
import os
import time
import traceback
import zipfile
from pathlib import Path
from typing import Any, Literal

from docx import Document
from docx.document import Document as DocumentType
from dotenv import load_dotenv
from pydantic import Field
from pydantic.fields import FieldInfo

from aworld.logs.util import Color
from examples.gaia.mcp_collections.base import ActionArguments, ActionCollection, ActionResponse
from examples.gaia.mcp_collections.documents.models import DocumentMetadata


class DOCXExtractionCollection(ActionCollection):
    """MCP service for DOCX/DOC document content extraction using python-docx.

    Supports extraction from DOCX and DOC files with comprehensive content parsing.
    Provides LLM-friendly text output with structured metadata and media file handling.
    """

    def __init__(self, arguments: ActionArguments) -> None:
        super().__init__(arguments)
        self._media_output_dir = self.workspace / "extracted_media"
        self._media_output_dir.mkdir(exist_ok=True)

        self.supported_extensions = {".docx", ".doc"}

        self._color_log("DOCX Extraction Service initialized", Color.green, "debug")
        self._color_log(f"Media output directory: {self._media_output_dir}", Color.blue, "debug")

    def _extract_images_from_docx(self, file_path: Path, file_stem: str) -> list[dict[str, str]]:
        """Extract embedded images from DOCX file.

        Args:
            file_path: Path to the DOCX file
            file_stem: Base name for saving files

        Returns:
            List of dictionaries containing image file paths and metadata
        """
        saved_media = []

        try:
            # DOCX files are ZIP archives
            with zipfile.ZipFile(file_path, "r") as docx_zip:
                # Look for media files in the word/media/ directory
                media_files = [f for f in docx_zip.namelist() if f.startswith("word/media/")]

                for idx, media_file in enumerate(media_files):
                    try:
                        # Extract file extension and create appropriate filename
                        original_name = Path(media_file).name
                        file_extension = Path(media_file).suffix

                        # Generate unique filename
                        media_filename = f"{file_stem}_media_{idx}_{original_name}"
                        media_path = self._media_output_dir / media_filename

                        # Extract and save the media file
                        with docx_zip.open(media_file) as source:
                            with open(media_path, "wb") as target:
                                target.write(source.read())

                        # Determine media type based on extension
                        media_type = "image"
                        if file_extension.lower() in [".mp3", ".wav", ".m4a", ".ogg"]:
                            media_type = "audio"
                        elif file_extension.lower() in [".mp4", ".avi", ".mov", ".wmv"]:
                            media_type = "video"
                        elif file_extension.lower() in [".png", ".jpg", ".jpeg", ".gif", ".bmp", ".tiff"]:
                            media_type = "image"
                        else:
                            media_type = "other"

                        saved_media.append(
                            {
                                "type": media_type,
                                "path": str(media_path),
                                "filename": media_filename,
                                "original_name": original_name,
                                "size_bytes": media_path.stat().st_size,
                            }
                        )

                        self._color_log(f"Extracted {media_type}: {media_filename}", Color.blue)

                    except Exception as e:
                        self.logger.error(f"Failed to extract media file {media_file}: {e}")

        except Exception as e:
            self.logger.warning(f"Could not extract media from DOCX: {e}")

        return saved_media

    def _extract_document_structure(self, doc: DocumentType) -> dict[str, Any]:
        """Extract document structure and metadata.

        Args:
            doc: python-docx Document object

        Returns:
            Dictionary containing document structure information
        """
        structure = {
            "paragraphs_count": len(doc.paragraphs),
            "tables_count": len(doc.tables),
            "sections_count": len(doc.sections),
            "styles_used": [],
            "has_headers_footers": False,
            "page_count": None,  # Not directly available in python-docx
        }

        # Collect unique styles used in the document
        styles_used = set()
        for paragraph in doc.paragraphs:
            if paragraph.style and paragraph.style.name:
                styles_used.add(paragraph.style.name)

        structure["styles_used"] = list(styles_used)

        # Check for headers and footers
        for section in doc.sections:
            if (section.header.paragraphs and any(p.text.strip() for p in section.header.paragraphs)) or (
                section.footer.paragraphs and any(p.text.strip() for p in section.footer.paragraphs)
            ):
                structure["has_headers_footers"] = True
                break

        return structure

    def _extract_tables_content(self, doc: DocumentType) -> list[dict[str, Any]]:
        """Extract content from all tables in the document.

        Args:
            doc: python-docx Document object

        Returns:
            List of dictionaries containing table data
        """
        tables_data = []

        for table_idx, table in enumerate(doc.tables):
            table_data = {
                "table_index": table_idx,
                "rows_count": len(table.rows),
                "columns_count": len(table.columns) if table.rows else 0,
                "data": [],
            }

            # Extract table content
            for _, row in enumerate(table.rows):
                row_data = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    row_data.append(cell_text)
                table_data["data"].append(row_data)

            tables_data.append(table_data)

        return tables_data

    def _convert_doc_to_docx(self, doc_path: Path) -> Path:
        """Convert .doc file to .docx using LibreOffice.

        Args:
            doc_path: Path to the .doc file

        Returns:
            Path to the converted .docx file
        """
        import subprocess

        output_dir = self._media_output_dir
        docx_path = output_dir / f"{doc_path.stem}.docx"

        # Check if already converted
        if docx_path.exists():
            self._color_log(f"Using existing converted file: {docx_path.name}", Color.blue)
            return docx_path

        self._color_log(f"Converting .doc to .docx: {doc_path.name}", Color.yellow)

        cmd = ["libreoffice", "--headless", "--convert-to", "docx", "--outdir", str(output_dir), str(doc_path)]

        try:
            result = subprocess.run(cmd, capture_output=True, text=True, check=True, timeout=60)  # pylint: disable=W0612
            if docx_path.exists():
                self._color_log(f"Conversion successful: {docx_path.name}", Color.green)
                return docx_path
            else:
                raise RuntimeError("Conversion completed but output file not found")

        except subprocess.CalledProcessError as e:
            self.logger.error(f"LibreOffice conversion failed: {e.stderr}")
            raise RuntimeError(f"Failed to convert .doc to .docx: {e.stderr}") from e
        except subprocess.TimeoutExpired as e:
            self.logger.error("LibreOffice conversion timed out")
            raise RuntimeError("Conversion timed out after 60 seconds") from e

    # Modify the _extract_content_from_docx method to handle .doc files
    def _extract_content_from_docx(
        self, file_path: Path, extract_tables: bool = True, extract_headers_footers: bool = True
    ) -> dict[str, Any]:
        """Extract content from DOCX file using python-docx.

        Args:
            file_path: Path to the DOCX/DOC file
            extract_tables: Whether to extract table content
            extract_headers_footers: Whether to extract headers and footers

        Returns:
            Dictionary containing extracted content and metadata
        """
        start_time = time.time()

        # Convert .doc to .docx if needed
        if file_path.suffix.lower() == ".doc":
            file_path = self._convert_doc_to_docx(file_path)

        try:
            # Load the document
            doc = Document(str(file_path))

            # Extract main text content
            paragraphs = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():  # Skip empty paragraphs
                    para_data = {"text": paragraph.text, "style": paragraph.style.name if paragraph.style else "Normal"}
                    paragraphs.append(para_data)

            # Extract document structure
            structure = self._extract_document_structure(doc)

            # Extract tables if requested
            tables = []
            if extract_tables:
                tables = self._extract_tables_content(doc)

            # Extract headers and footers if requested
            headers_footers = []
            if extract_headers_footers:
                for section_idx, section in enumerate(doc.sections):
                    # Extract header content
                    header_text = []
                    for para in section.header.paragraphs:
                        if para.text.strip():
                            header_text.append(para.text)

                    # Extract footer content
                    footer_text = []
                    for para in section.footer.paragraphs:
                        if para.text.strip():
                            footer_text.append(para.text)

                    if header_text or footer_text:
                        headers_footers.append(
                            {"section_index": section_idx, "header": header_text, "footer": footer_text}
                        )

            processing_time = time.time() - start_time

            return {
                "paragraphs": paragraphs,
                "tables": tables,
                "headers_footers": headers_footers,
                "structure": structure,
                "processing_time": processing_time,
                "word_count": sum(len(p["text"].split()) for p in paragraphs),
                "character_count": sum(len(p["text"]) for p in paragraphs),
            }

        except Exception as e:
            self.logger.error(f"Failed to extract content from DOCX: {e}")
            raise

    def _format_content_for_llm(
        self, extraction_result: dict[str, Any], output_format: str, include_structure: bool = True
    ) -> str:
        """Format extracted DOCX content to be LLM-friendly.

        Args:
            extraction_result: Dictionary containing extracted content
            output_format: Desired output format
            include_structure: Whether to include document structure information

        Returns:
            Formatted content string
        """
        if output_format.lower() == "markdown":
            content_parts = []

            # Add document structure info if requested
            if include_structure:
                structure = extraction_result["structure"]
                content_parts.append("# Document Structure\n")
                content_parts.append(f"- **Paragraphs**: {structure['paragraphs_count']}\n")
                content_parts.append(f"- **Tables**: {structure['tables_count']}\n")
                content_parts.append(f"- **Sections**: {structure['sections_count']}\n")
                content_parts.append(f"- **Word Count**: {extraction_result['word_count']}\n")
                content_parts.append(f"- **Character Count**: {extraction_result['character_count']}\n")
                if structure["styles_used"]:
                    content_parts.append(f"- **Styles Used**: {', '.join(structure['styles_used'])}\n")
                content_parts.append("\n---\n\n")

            # Add main content
            content_parts.append("# Document Content\n\n")

            # Add paragraphs
            for para in extraction_result["paragraphs"]:
                # Format based on style
                text = para["text"]
                style = para["style"]

                if "Heading" in style:
                    # Convert heading styles to markdown headers
                    if "Heading 1" in style:
                        content_parts.append(f"# {text}\n\n")
                    elif "Heading 2" in style:
                        content_parts.append(f"## {text}\n\n")
                    elif "Heading 3" in style:
                        content_parts.append(f"### {text}\n\n")
                    else:
                        content_parts.append(f"#### {text}\n\n")
                else:
                    content_parts.append(f"{text}\n\n")

            # Add tables
            if extraction_result["tables"]:
                content_parts.append("\n## Tables\n\n")
                for table_idx, table in enumerate(extraction_result["tables"]):
                    content_parts.append(f"### Table {table_idx + 1}\n\n")

                    if table["data"]:
                        # Create markdown table
                        headers = table["data"][0] if table["data"] else []
                        if headers:
                            content_parts.append("| " + " | ".join(headers) + " |\n")
                            content_parts.append("|" + "---|" * len(headers) + "\n")

                            for row in table["data"][1:]:
                                content_parts.append("| " + " | ".join(row) + " |\n")
                        content_parts.append("\n")

            # Add headers and footers
            if extraction_result["headers_footers"]:
                content_parts.append("\n## Headers and Footers\n\n")
                for hf in extraction_result["headers_footers"]:
                    if hf["header"]:
                        content_parts.append(f"**Header (Section {hf['section_index'] + 1}):**\n")
                        for header_line in hf["header"]:
                            content_parts.append(f"{header_line}\n")
                        content_parts.append("\n")

                    if hf["footer"]:
                        content_parts.append(f"**Footer (Section {hf['section_index'] + 1}):**\n")
                        for footer_line in hf["footer"]:
                            content_parts.append(f"{footer_line}\n")
                        content_parts.append("\n")

            return "".join(content_parts)

        elif output_format.lower() == "json":
            # Return structured JSON
            return json.dumps(extraction_result, indent=2, ensure_ascii=False)

        elif output_format.lower() == "html":
            # Convert to HTML
            html_parts = ["<html><body>"]

            if include_structure:
                html_parts.append("<h1>Document Structure</h1>")
                structure = extraction_result["structure"]
                html_parts.append(f"<p><strong>Paragraphs:</strong> {structure['paragraphs_count']}</p>")
                html_parts.append(f"<p><strong>Tables:</strong> {structure['tables_count']}</p>")
                html_parts.append(f"<p><strong>Word Count:</strong> {extraction_result['word_count']}</p>")
                html_parts.append("<hr>")

            html_parts.append("<h1>Document Content</h1>")

            # Add paragraphs
            for para in extraction_result["paragraphs"]:
                text = para["text"].replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
                style = para["style"]

                if "Heading" in style:
                    if "Heading 1" in style:
                        html_parts.append(f"<h1>{text}</h1>")
                    elif "Heading 2" in style:
                        html_parts.append(f"<h2>{text}</h2>")
                    elif "Heading 3" in style:
                        html_parts.append(f"<h3>{text}</h3>")
                    else:
                        html_parts.append(f"<h4>{text}</h4>")
                else:
                    html_parts.append(f"<p>{text}</p>")

            # Add tables
            if extraction_result["tables"]:
                html_parts.append("<h2>Tables</h2>")
                for table_idx, table in enumerate(extraction_result["tables"]):
                    html_parts.append(f"<h3>Table {table_idx + 1}</h3>")
                    html_parts.append("<table border='1'>")

                    for row_idx, row in enumerate(table["data"]):
                        html_parts.append("<tr>")
                        tag = "th" if row_idx == 0 else "td"
                        for cell in row:
                            cell_text = cell.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
                            html_parts.append(f"<{tag}>{cell_text}</{tag}>")
                        html_parts.append("</tr>")

                    html_parts.append("</table>")

            html_parts.append("</body></html>")
            return "".join(html_parts)

        else:
            # Plain text format
            text_parts = []

            # Add main content
            for para in extraction_result["paragraphs"]:
                text_parts.append(para["text"])

            # Add tables as plain text
            if extraction_result["tables"]:
                text_parts.append("\n\n=== TABLES ===\n")
                for table_idx, table in enumerate(extraction_result["tables"]):
                    text_parts.append(f"\nTable {table_idx + 1}:\n")
                    for row in table["data"]:
                        text_parts.append("\t".join(row) + "\n")

            return "\n\n".join(text_parts)

    def mcp_extract_docx_content(
        self,
        file_path: str = Field(description="Path to the DOCX/DOC document file to extract content from"),
        output_format: Literal["markdown", "json", "html", "text"] = Field(
            default="markdown", description="Output format: 'markdown', 'json', 'html', or 'text'"
        ),
        extract_images: bool = Field(default=True, description="Whether to extract and save embedded images and media"),
        extract_tables: bool = Field(default=True, description="Whether to extract table content"),
        extract_headers_footers: bool = Field(default=True, description="Whether to extract headers and footers"),
        include_structure: bool = Field(
            default=True, description="Whether to include document structure information in output"
        ),
    ) -> ActionResponse:
        """Extract content from DOCX/DOC documents using python-docx.

        This tool provides comprehensive DOCX/DOC document content extraction with support for:
        - DOCX and DOC files
        - Text extraction with style preservation
        - Table extraction and formatting
        - Headers and footers extraction
        - Embedded media extraction (images, audio, etc.)
        - Multiple output formats (Markdown, JSON, HTML, Text)
        - Document structure analysis

        Args:
            file_path: Path to the DOCX/DOC file
            output_format: Desired output format
            extract_images: Extract embedded media files
            extract_tables: Extract table content
            extract_headers_footers: Extract headers and footers
            include_structure: Include document structure info

        Returns:
            ActionResponse with extracted content, metadata, and media file paths
        """
        try:
            # Handle FieldInfo objects from pydantic
            if isinstance(file_path, FieldInfo):
                file_path = file_path.default
            if isinstance(output_format, FieldInfo):
                output_format = output_format.default
            if isinstance(extract_images, FieldInfo):
                extract_images = extract_images.default
            if isinstance(extract_tables, FieldInfo):
                extract_tables = extract_tables.default
            if isinstance(extract_headers_footers, FieldInfo):
                extract_headers_footers = extract_headers_footers.default
            if isinstance(include_structure, FieldInfo):
                include_structure = include_structure.default

            # Validate input file
            file_path: Path = self._validate_file_path(file_path)
            self._color_log(f"Processing DOCX document: {file_path.name}", Color.cyan)

            # Extract embedded media if requested
            saved_media = []
            if extract_images and file_path.suffix.lower() == ".docx":
                saved_media = self._extract_images_from_docx(file_path, file_path.stem)

            # Extract document content
            extraction_result = self._extract_content_from_docx(
                file_path, extract_tables=extract_tables, extract_headers_footers=extract_headers_footers
            )

            # Format content for LLM consumption
            formatted_content = self._format_content_for_llm(
                extraction_result, output_format, include_structure=include_structure
            )

            # Prepare metadata
            file_stats = file_path.stat()
            document_metadata = DocumentMetadata(
                file_name=file_path.name,
                file_size=file_stats.st_size,
                file_type=file_path.suffix.lower(),
                absolute_path=str(file_path.absolute()),
                page_count=None,  # Not directly available for DOCX
                processing_time=extraction_result["processing_time"],
                extracted_images=[media["path"] for media in saved_media if media["type"] == "image"],
                extracted_media=saved_media,
                output_format=output_format,
                llm_enhanced=False,
                ocr_applied=False,
            )

            # Add DOCX-specific metadata
            docx_metadata = {
                "paragraphs_count": extraction_result["structure"]["paragraphs_count"],
                "tables_count": extraction_result["structure"]["tables_count"],
                "sections_count": extraction_result["structure"]["sections_count"],
                "word_count": extraction_result["word_count"],
                "character_count": extraction_result["character_count"],
                "styles_used": extraction_result["structure"]["styles_used"],
                "has_headers_footers": extraction_result["structure"]["has_headers_footers"],
                "has_embedded_media": len(saved_media) > 0,
                "media_files_count": len(saved_media),
            }

            # Merge metadata
            final_metadata = {**document_metadata.model_dump(), **docx_metadata}

            self._color_log(
                f"Successfully extracted DOCX content from {file_path.name} "
                f"({extraction_result['word_count']} words, {extraction_result['structure']['tables_count']} tables, "
                f"{len(saved_media)} media files)",
                Color.green,
            )

            return ActionResponse(success=True, message=formatted_content, metadata=final_metadata)

        except FileNotFoundError as e:
            self.logger.error(f"File not found: {str(e)}")
            return ActionResponse(
                success=False, message=f"File not found: {str(e)}", metadata={"error_type": "file_not_found"}
            )
        except ValueError as e:
            self.logger.error(f"Invalid input: {str(e)}")
            return ActionResponse(
                success=False,
                message=f"Invalid input: {str(e)}",
                metadata={"error_type": "invalid_input"},
            )
        except ImportError as e:
            self.logger.error(f"Missing dependency: {str(e)}")
            return ActionResponse(
                success=False,
                message=f"Missing dependency: {str(e)}. Please install python-docx: pip install python-docx",
                metadata={"error_type": "missing_dependency"},
            )
        except Exception as e:
            self.logger.error(f"DOCX extraction failed: {str(e)}: {traceback.format_exc()}")
            return ActionResponse(
                success=False,
                message=f"DOCX extraction failed: {str(e)}",
                metadata={"error_type": "extraction_error"},
            )

    def mcp_list_supported_formats(self) -> ActionResponse:
        """List all supported document formats for extraction.

        Returns:
            ActionResponse with list of supported file formats and their descriptions
        """
        supported_formats = {
            "DOCX": "Microsoft Word Open XML Document (.docx)",
            "DOC": "Microsoft Word Document (.doc) - limited support",
        }

        format_list = "\n".join(
            [f"**{format_name}**: {description}" for format_name, description in supported_formats.items()]
        )

        return ActionResponse(
            success=True,
            message=f"Supported document formats:\n\n{format_list}\n\n"
            "**Note**: DOC format support is limited. For best results, convert to DOCX format.",
            metadata={"supported_formats": list(supported_formats.keys()), "total_formats": len(supported_formats)},
        )


# Example usage and entry point
if __name__ == "__main__":
    load_dotenv()

    # Default arguments for testing
    args = ActionArguments(
        name="docx_extraction_service",
        transport="stdio",
        workspace=os.getenv("AWORLD_WORKSPACE", "~"),
    )

    # Initialize and run the DOCX extraction service
    try:
        service = DOCXExtractionCollection(args)
        service.run()
    except Exception as e:
        print(f"An error occurred: {e}: {traceback.format_exc()}")
